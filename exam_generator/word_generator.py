"""Word生成模块，用于将试题数据转换为Word文档"""

# 延迟导入，减少模块加载时间


def set_page_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.25):
    """
    设置文档页边距
    
    参数:
        doc: Document对象
        top: 上边距（英寸）
        bottom: 下边距（英寸）
        left: 左边距（英寸）
        right: 右边距（英寸）
    """
    # 在函数内部导入，实现延迟加载
    from docx.shared import Inches
    
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def set_header_footer(doc, header_text="", footer_text=""):
    """
    设置文档页眉页脚
    
    参数:
        doc: Document对象
        header_text: 页眉文本
        footer_text: 页脚文本
    """
    # 在函数内部导入，实现延迟加载
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    for section in doc.sections:
        # 设置页眉
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = header_text
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 设置页脚
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = footer_text
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate_exam_document(questions, output_file, title="考试试卷", shuffle_options=False):
    """
    生成考试试卷Word文档
    
    参数:
        questions: 试题对象列表
        output_file: 输出文件路径
        title: 试卷标题
        shuffle_options: 是否打乱选项
    """
    # 在函数内部导入，实现延迟加载
    from docx import Document
    
    try:
        # 创建新文档
        doc = Document()
        
        # 添加试卷名称
        if questions:
            start_id = 1
            end_id = len(questions)
            doc.add_paragraph(f"试卷名称：{title}（{start_id}-{end_id}题）")
        else:
            doc.add_paragraph(f"试卷名称：{title}")
        
        # 先对所有需要打乱选项的试题进行处理，避免重复打乱
        if shuffle_options:
            for question in questions:
                question.shuffle_options()
        
        # 按题型分组处理
        question_types = {}
        for q in questions:
            if q.question_type not in question_types:
                question_types[q.question_type] = []
            question_types[q.question_type].append(q)
        
        # 处理每种题型
        for q_type, q_list in question_types.items():
            # 获取该题型的分值（假设同一题型分值相同）
            score = q_list[0].score if q_list else 0
            
            # 特殊处理阅读理解题型
            if q_type == '阅读理解':
                # 将阅读理解题目按文章分组（每篇文章后有5道题）
                reading_groups = []
                current_group = []
                
                for question in q_list:
                    if not current_group:
                        # 开始新组
                        current_group = [question]
                    elif len(current_group) < 5:
                        # 如果当前组不足5道题，继续添加
                        current_group.append(question)
                    else:
                        # 当前组已满5道题，保存并开始新组
                        reading_groups.append(current_group)
                        current_group = [question]
                
                # 确保添加最后一组
                if current_group:
                    reading_groups.append(current_group)
                
                # 处理每个阅读理解组
                for group_idx, group in enumerate(reading_groups):
                    # 添加题型说明
                    doc.add_paragraph(f"<TYPE.TAG>文本行")
                    doc.add_paragraph(f"{q_type},每小题{score}分。")
                    doc.add_paragraph(f"<TYPE.TAG>文本行")
                    doc.add_paragraph()  # 添加空行
                    
                    # 添加阅读理解文章（只在每组开始显示一次）
                    passage_para = doc.add_paragraph(group[0].reading_passage)
                    passage_para.style = 'BodyText'
                    doc.add_paragraph()  # 添加空行
                    
                    # 处理该组的5道题目
                    for q_idx, question in enumerate(group):
                        # 添加试题编号和题目内容，使用连续序号
                        doc.add_paragraph(f"<TYPE.TAG>单项选择题")
                        doc.add_paragraph()  # 添加空行
                        doc.add_paragraph(f"{group_idx * 5 + q_idx + 1}.{question.title}")
                        
                        # 添加选项
                        for option_label, option_text in [
                            ("A", question.option_a),
                            ("B", question.option_b),
                            ("C", question.option_c),
                            ("D", question.option_d)
                        ]:
                            if option_text and str(option_text).strip():
                                doc.add_paragraph(f"{option_label}.{option_text}")
                        
                        # 添加答案和分数
                        doc.add_paragraph(f"参考答案:{question.correct_option}")
                        doc.add_paragraph(f"分数:{question.score}")
                        
                        # 添加解析
                        analysis_text = str(question.analysis).strip() if question.analysis else ""
                        doc.add_paragraph(f"解析:{analysis_text}")
                        doc.add_paragraph()  # 添加空行
            else:
                # 处理其他题型，添加题型说明
                display_q_type = "单项选择题" if q_type == "选择题" else q_type
                doc.add_paragraph(f"<TYPE.TAG>{display_q_type}")
                doc.add_paragraph(f"{q_type},每小题{score}分。")
                doc.add_paragraph(f"<TYPE.TAG>{display_q_type}")
                doc.add_paragraph()  # 添加空行
                
                # 处理其他题型，使用连续序号
                for idx, question in enumerate(q_list, 1):
                    # 添加试题编号和题目内容，使用连续序号
                    display_q_type = "单项选择题" if q_type == "选择题" else q_type
                    doc.add_paragraph(f"<TYPE.TAG>{display_q_type}")
                    doc.add_paragraph(f"{idx}.{question.title}")
                    
                    # 添加选项
                    for option_label, option_text in [
                        ("A", question.option_a),
                        ("B", question.option_b),
                        ("C", question.option_c),
                        ("D", question.option_d)
                    ]:
                        if option_text and str(option_text).strip():
                            doc.add_paragraph(f"{option_label}.{option_text}")
                    
                    # 添加答案和分数
                    doc.add_paragraph(f"参考答案:{question.correct_option}")
                    doc.add_paragraph(f"分数:{question.score}")
                    
                    # 添加解析（确保每道题都有解析部分）
                    analysis_text = str(question.analysis).strip() if question.analysis else ""
                    doc.add_paragraph(f"解析:{analysis_text}")
                    doc.add_paragraph()  # 添加空行
        
        # 保存文档
        doc.save(output_file)
        
        return True
        
    except Exception as e:
        print(f"生成Word文档时出错: {str(e)}")
        return False